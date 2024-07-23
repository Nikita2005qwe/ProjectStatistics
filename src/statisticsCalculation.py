from functionOfUploadDict import uploadDict
from docx import Document
from PyQt5.QtWidgets import QProgressBar

class StatisticsCalculation:
    """
    Класс с помощью которого можно посчитать число слов и букв в тексте
    методы возвращают словари со статистикой по количеству слов и букв
    """
    def __init__(self, path: str, progress_bar: QProgressBar):
        self.numberOfWords: dict[str, int] = self.wordCount(path, progress_bar)
        self.numberOfLetters: dict[str, int] = self.letterCount(self.numberOfWords, progress_bar)

    @staticmethod
    def wordCount(path: str, progress_bar: QProgressBar) -> dict[str, int]:
        number_of_words: dict[str, int] = dict()
        if ".docx" in path or ".doc" in path:
            doc: Document = Document(path)
            file_size = len(doc.paragraphs)
            cnt = 0
            for paragraph in doc.paragraphs:
                uploadDict(paragraph.text, number_of_words, lambda: 1, True)
                progress_bar.setValue(int(cnt / file_size * 50))
                cnt += 1
            progress_bar.setValue(int(50))
            return number_of_words
            
        file_size = sum(1 for _ in open(path, "r", encoding="utf-8"))
        with open(path, encoding="utf-8") as file:
            cnt = 0
            for string in file:
                uploadDict(string, number_of_words, lambda: 1, True)
                progress_bar.setValue(int(cnt/file_size*50))
                cnt += 1
            progress_bar.setValue(int(50))
        return number_of_words
    
    @staticmethod
    def letterCount(number_of_words: dict[str, int], progress_bar: QProgressBar) -> dict[str, int]:
        number_of_letters: dict[str, int] = dict()
        file_size = len(number_of_words)
        cnt = 0
        for word in number_of_words:
            uploadDict(word, number_of_letters, lambda: number_of_words[word], False)
            progress_bar.setValue(int(50+(cnt / file_size) * 50))
            cnt += 1
        progress_bar.setValue(int(100))
        return number_of_letters
    
    def getNumberOfWords(self) -> dict[str, int]:
        return self.numberOfWords
    
    def getNumberOfLetters(self) -> dict[str, int]:
        return self.numberOfLetters
